from __future__ import annotations

import asyncio
import hashlib
import os
from collections.abc import Awaitable, Callable, Iterator, Mapping, Sequence
from contextlib import suppress
from datetime import UTC, datetime, timedelta
from io import BytesIO
from pathlib import Path
from typing import Any, cast
from uuid import UUID, uuid4

import httpx
import pytest
from alembic import command
from alembic.config import Config
from docx import Document
from pydantic import SecretStr
from sqlalchemy import delete, select, update
from sqlalchemy.engine import make_url
from sqlalchemy.exc import IntegrityError

from mongars.config import Environment, Settings
from mongars.db.models import (
    DocumentStaging,
    EpisodicEvent,
    MemoryChunk,
    MemoryDocument,
    MemoryDocumentProvenance,
    TaskQueue,
)
from mongars.db.session import Database
from mongars.embeddings.models import EmbeddingBatch
from mongars.embeddings.service import EmbeddingService
from mongars.inference import ChatMessage, ChatResponse, HealthStatus, JsonValue
from mongars.ingestion.errors import ParserTimeoutError
from mongars.ingestion.isolation import DocumentParser, ParserHealth
from mongars.ingestion.models import ExtractedContent, ValidatedUpload
from mongars.ingestion.service import DocumentIngestionService
from mongars.ingestion.staging import DocumentStagingRepository
from mongars.main import create_app
from mongars.rm.repository import TaskRepository
from mongars.rm.worker import Worker

_RAW_DATABASE_URL = os.getenv("MONGARS_TEST_DATABASE_URL", "").strip()
if not _RAW_DATABASE_URL:
    pytest.skip(
        "MONGARS_TEST_DATABASE_URL is required for PostgreSQL integration tests",
        allow_module_level=True,
    )

pytestmark = pytest.mark.integration


def _psycopg_url(value: str) -> str:
    url = make_url(value)
    if url.get_backend_name() != "postgresql":
        raise ValueError("MONGARS_TEST_DATABASE_URL must target PostgreSQL")
    return url.set(drivername="postgresql+psycopg").render_as_string(hide_password=False)


DATABASE_URL = _psycopg_url(_RAW_DATABASE_URL)


@pytest.fixture(scope="module", autouse=True)
def migrated_database() -> Iterator[None]:
    root = Path(__file__).resolve().parents[2]
    config = Config(str(root / "alembic.ini"))
    previous_url = os.environ.get("MONGARS_DATABASE_URL")
    os.environ["MONGARS_DATABASE_URL"] = DATABASE_URL
    try:
        command.upgrade(config, "head")
        yield
    finally:
        if previous_url is None:
            os.environ.pop("MONGARS_DATABASE_URL", None)
        else:
            os.environ["MONGARS_DATABASE_URL"] = previous_url


class DeterministicInference:
    async def chat(
        self,
        messages: Sequence[ChatMessage],
        *,
        model: str | None = None,
        options: Mapping[str, JsonValue] | None = None,
    ) -> ChatResponse:
        del messages, options
        return ChatResponse(content="unused", model=model or "deterministic-chat")

    async def health(self) -> HealthStatus:
        return HealthStatus(
            backend="deterministic",
            backend_reachable=True,
            chat_model_ready=True,
            embedding_model_ready=True,
            latency_ms=0.0,
        )

    async def aclose(self) -> None:
        return None


type Probe = Callable[[], Awaitable[None]]


class ProbedEmbeddingProvider:
    provider_name = "deterministic"
    model_name = "nomic-embed-text"

    def __init__(self, *, probe: Probe | None = None) -> None:
        self.calls = 0
        self.probe = probe

    async def resolve_model_digest(self) -> str:
        return "0a109f422b47e3a30ba2b10eca18548e944e8a23073ee3f3e947efcf3c45e59f"

    async def embed(
        self,
        texts: Sequence[str],
        *,
        expected_dimension: int,
    ) -> EmbeddingBatch:
        self.calls += 1
        if self.probe is not None:
            await self.probe()
        vector = (1.0, *([0.0] * (expected_dimension - 1)))
        return EmbeddingBatch(
            embeddings=tuple(vector for _text in texts),
            model=self.model_name,
            model_digest=await self.resolve_model_digest(),
            dimension=expected_dimension,
            latency_ms=0.0,
        )

    async def aclose(self) -> None:
        return None


class ProbedDocumentParser:
    def __init__(self, *, probe: Probe | None = None) -> None:
        self.calls = 0
        self.probe = probe

    async def extract(
        self,
        upload: ValidatedUpload,
    ) -> ExtractedContent:
        self.calls += 1
        if self.probe is not None:
            await self.probe()
        return DocumentIngestionService().extract_content(upload)

    async def aclose(self) -> None:
        return None

    async def health(self) -> ParserHealth:
        return ParserHealth(healthy=True, parser_version="integration-parser-v1")


class FlakyDocumentParser(ProbedDocumentParser):
    async def extract(self, upload: ValidatedUpload) -> ExtractedContent:
        self.calls += 1
        if self.calls == 1:
            raise ParserTimeoutError("deterministic parser timeout")
        if self.probe is not None:
            await self.probe()
        return DocumentIngestionService().extract_content(upload)


class GatedDocumentParser(ProbedDocumentParser):
    def __init__(self) -> None:
        super().__init__()
        self.started = asyncio.Event()
        self.release = asyncio.Event()

    async def extract(self, upload: ValidatedUpload) -> ExtractedContent:
        self.calls += 1
        self.started.set()
        await self.release.wait()
        return DocumentIngestionService().extract_content(upload)


def _settings(owner_id: str, token: str) -> Settings:
    return Settings(
        environment=Environment.TEST,
        owner_id=owner_id,
        api_token=SecretStr(token),
        approval_hmac_key=SecretStr("document-integration-approval-key"),
        database_url=DATABASE_URL,
        web_search_enabled=False,
        memory_chunk_tokens=64,
        memory_chunk_overlap_tokens=8,
        worker_lease_seconds=10,
        retention_sweep_seconds=10,
    )


def _embedding_service(provider: ProbedEmbeddingProvider) -> EmbeddingService:
    return EmbeddingService(
        provider=provider,
        expected_dimension=768,
        batch_size=16,
    )


async def _clean_owner(database: Database, owner_id: str) -> None:
    async with database.session_factory() as session, session.begin():
        await session.execute(delete(EpisodicEvent).where(EpisodicEvent.owner_id == owner_id))
        await session.execute(delete(TaskQueue).where(TaskQueue.owner_id == owner_id))
        await session.execute(delete(MemoryDocument).where(MemoryDocument.owner_id == owner_id))


async def _upload_document(
    client: httpx.AsyncClient,
    *,
    headers: dict[str, str],
    content: bytes,
    filename: str,
    mime_type: str,
) -> dict[str, Any]:
    response = await client.post(
        "/v1/documents",
        headers=headers,
        files={"file": (filename, content, mime_type)},
        data={
            "declared_size": str(len(content)),
            "source_timestamp": "2026-07-22T08:30:00-04:00",
            "title": "Reviewed TXT upload",
            "sensitivity": "restricted",
            "retention_class": "ttl_30d",
        },
    )
    assert response.status_code == 202, response.text
    return cast(dict[str, Any], response.json())


async def _upload_txt(
    client: httpx.AsyncClient,
    *,
    headers: dict[str, str],
    content: bytes,
    filename: str = "notes.txt",
) -> dict[str, Any]:
    return await _upload_document(
        client,
        headers=headers,
        content=content,
        filename=filename,
        mime_type="text/plain",
    )


def _make_pdf(text: str) -> bytes:
    objects: list[bytes] = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>"
        ),
    ]
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode()
    objects.extend(
        (
            b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
            b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        )
    )
    output = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for index, value in enumerate(objects, 1):
        offsets.append(len(output))
        output.extend(f"{index} 0 obj\n".encode())
        output.extend(value)
        output.extend(b"\nendobj\n")
    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets:
        output.extend(f"{offset:010d} 00000 n \n".encode())
    output.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode()
    )
    return bytes(output)


def _make_docx(text: str) -> bytes:
    document = Document()
    document.add_heading("Durable document", level=1)
    document.add_paragraph(text)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


@pytest.mark.asyncio
async def test_multipart_upload_approval_worker_persists_provenance_without_long_tx() -> None:
    owner_id = f"document-runtime-{uuid4().hex}"
    token = uuid4().hex
    settings = _settings(owner_id, token)
    database = Database(settings)
    inference = DeterministicInference()
    provider = ProbedEmbeddingProvider()
    embeddings = _embedding_service(provider)
    parser = ProbedDocumentParser()
    application = create_app(
        settings=settings,
        database=database,
        inference=inference,
        embeddings=embeddings,
    )
    headers = {"Authorization": f"Bearer {token}"}
    content = f"Approved document marker {uuid4().hex}.".encode()

    try:
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=application),
            base_url="http://testserver",
        ) as client:
            request_started_at = datetime.now(UTC)
            created = await _upload_txt(client, headers=headers, content=content)
            request_completed_at = datetime.now(UTC)
            assert set(created) == {
                "id",
                "kind",
                "status",
                "risk_level",
                "action_digest",
            }
            assert created["kind"] == "document.ingest"
            assert created["status"] == "waiting_approval"
            assert created["risk_level"] == "local_mutation"
            assert len(created["action_digest"]) == 64
            task_id = UUID(created["id"])

            async with database.session_factory() as session:
                task = await session.get(TaskQueue, task_id)
                assert task is not None
                received_at = datetime.fromisoformat(task.payload["received_at"])
                assert request_started_at <= received_at <= request_completed_at
                assert task.payload["source_time_basis"] == "user_supplied"
                assert task.payload["source_timestamp"] == "2026-07-22T12:30:00Z"
                staging_id = UUID(task.payload["staging_id"])
                staged = await session.get(DocumentStaging, staging_id)
                assert staged is not None
                assert staged.owner_id == owner_id
                assert staged.task_id == task_id
                assert staged.content == content
                assert staged.source_sha256 == hashlib.sha256(content).digest()
                assert staged.received_at == received_at
                assert staged.source_timestamp == datetime(2026, 7, 22, 12, 30, tzinfo=UTC)
                assert staged.expires_at > received_at + timedelta(
                    seconds=settings.document_staging_ttl_seconds
                )
                assert staged.expires_at <= request_completed_at + timedelta(
                    seconds=settings.document_staging_ttl_seconds
                )
                assert (
                    await DocumentStagingRepository(session).get_for_owner(
                        staging_id=staging_id,
                        owner_id=f"foreign-{owner_id}",
                    )
                    is None
                )

            async def probe_no_worker_row_locks() -> None:
                async with database.session_factory() as session, session.begin():
                    locked_task = await session.scalar(
                        select(TaskQueue)
                        .where(TaskQueue.id == task_id)
                        .with_for_update(nowait=True)
                    )
                    locked_stage = await session.scalar(
                        select(DocumentStaging)
                        .where(DocumentStaging.id == staging_id)
                        .with_for_update(nowait=True)
                    )
                    assert locked_task is not None
                    assert locked_stage is not None

            parser.probe = probe_no_worker_row_locks
            provider.probe = probe_no_worker_row_locks

            approved = await client.post(
                f"/v1/tasks/{task_id}/approve",
                headers=headers,
                json={"action_digest": created["action_digest"]},
            )
            assert approved.status_code == 200, approved.text
            assert approved.json()["status"] == "queued"

            worker = Worker(
                settings=settings,
                database=database,
                inference=inference,
                embeddings=embeddings,
                document_parser=cast(DocumentParser, parser),
            )
            assert await worker.run_once() is True

            detail = await client.get(f"/v1/tasks/{task_id}", headers=headers)
            assert detail.status_code == 200
            task_result = detail.json()
            assert task_result["status"] == "done"
            assert task_result["attempt_count"] == 1
            assert task_result["result"]["created"] is True
            assert task_result["result"]["chunk_count"] == 1
            assert task_result["result"]["provenance"] == {
                "sha256": hashlib.sha256(content).hexdigest(),
                "original_filename": "notes.txt",
                "validated_mime_type": "text/plain",
                "byte_size": len(content),
                "extracted_character_count": len(content.decode()),
                "page_count": None,
                "section_count": 1,
                "parser_name": "utf8-text",
                "parser_version": "1",
                "ingestion_task_id": str(task_id),
                "owner_id": owner_id,
                "sensitivity": "restricted",
                "retention_class": "ttl_30d",
                "source_timestamp": "2026-07-22T12:30:00+00:00",
                "received_at": received_at.isoformat(),
                "source_time_basis": "user_supplied",
            }
            document_id = UUID(task_result["result"]["document_id"])

            async with database.session_factory() as session:
                document = await session.get(MemoryDocument, document_id)
                assert document is not None
                assert document.owner_id == owner_id
                assert document.source_sha256 == hashlib.sha256(content).digest()
                assert document.sensitivity == "restricted"
                assert document.retention_class == "ttl_30d"
                assert document.mime_type == "text/plain"
                assert document.metadata_json == task_result["result"]["provenance"]
                assert await session.get(DocumentStaging, staging_id) is None
                provenances = list(
                    (
                        await session.scalars(
                            select(MemoryDocumentProvenance).where(
                                MemoryDocumentProvenance.document_id == document_id
                            )
                        )
                    ).all()
                )
                assert len(provenances) == 1
                assert provenances[0].metadata_json == task_result["result"]["provenance"]
                event_types = list(
                    (
                        await session.scalars(
                            select(EpisodicEvent.event_type).where(
                                EpisodicEvent.owner_id == owner_id,
                                EpisodicEvent.trace_id == task_result["trace_id"],
                            )
                        )
                    ).all()
                )
                assert "document_ingested" in event_types
                assert "task_completed" in event_types

            assert parser.calls == 1
            assert provider.calls == 1
    finally:
        await _clean_owner(database, owner_id)
        await embeddings.aclose()
        await inference.aclose()
        await database.close()


@pytest.mark.parametrize(
    ("filename", "mime_type", "content", "marker", "parser_name"),
    [
        (
            "architecture.md",
            "text/markdown",
            b"# Main ingestion\n\nDurable Markdown marker.",
            "Durable Markdown marker.",
            "markdown-plaintext",
        ),
        (
            "architecture.html",
            "text/html",
            (
                b"<!doctype html><html><body><main><h1>Main ingestion</h1>"
                b"<p>Durable HTML marker.</p></main></body></html>"
            ),
            "Durable HTML marker.",
            "beautifulsoup-html",
        ),
        (
            "architecture.pdf",
            "application/pdf",
            _make_pdf("Durable PDF marker."),
            "Durable PDF marker.",
            "pypdf",
        ),
        (
            "architecture.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            _make_docx("Durable DOCX marker."),
            "Durable DOCX marker.",
            "python-docx",
        ),
    ],
    ids=("markdown", "html", "pdf", "docx"),
)
@pytest.mark.asyncio
async def test_supported_format_traverses_approved_durable_ingestion(
    filename: str,
    mime_type: str,
    content: bytes,
    marker: str,
    parser_name: str,
) -> None:
    owner_id = f"document-format-{uuid4().hex}"
    token = uuid4().hex
    settings = _settings(owner_id, token)
    database = Database(settings)
    inference = DeterministicInference()
    provider = ProbedEmbeddingProvider()
    embeddings = _embedding_service(provider)
    parser = ProbedDocumentParser()
    application = create_app(
        settings=settings,
        database=database,
        inference=inference,
        embeddings=embeddings,
    )
    headers = {"Authorization": f"Bearer {token}"}

    try:
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=application),
            base_url="http://testserver",
        ) as client:
            created = await _upload_document(
                client,
                headers=headers,
                content=content,
                filename=filename,
                mime_type=mime_type,
            )
            task_id = UUID(created["id"])
            async with database.session_factory() as session:
                task = await session.get(TaskQueue, task_id)
                assert task is not None
                staging_id = UUID(task.payload["staging_id"])

            approved = await client.post(
                f"/v1/tasks/{task_id}/approve",
                headers=headers,
                json={"action_digest": created["action_digest"]},
            )
            assert approved.status_code == 200, approved.text

            worker = Worker(
                settings=settings,
                database=database,
                inference=inference,
                embeddings=embeddings,
                document_parser=cast(DocumentParser, parser),
            )
            assert await worker.run_once() is True

            async with database.session_factory() as session:
                task = await session.get(TaskQueue, task_id)
                assert task is not None
                assert task.status == "done"
                assert task.result is not None
                document_id = UUID(task.result["document_id"])
                assert task.result["provenance"]["validated_mime_type"] == mime_type
                assert task.result["provenance"]["parser_name"] == parser_name
                assert task.result["provenance"]["original_filename"] == filename
                assert task.result["provenance"]["sha256"] == hashlib.sha256(content).hexdigest()
                assert await session.get(DocumentStaging, staging_id) is None
                chunks = list(
                    (
                        await session.scalars(
                            select(MemoryChunk).where(MemoryChunk.document_id == document_id)
                        )
                    ).all()
                )
                assert chunks
                assert marker in "\n".join(chunk.plaintext for chunk in chunks)
                assert (
                    await session.scalar(
                        select(EpisodicEvent.id).where(
                            EpisodicEvent.owner_id == owner_id,
                            EpisodicEvent.trace_id == task.trace_id,
                            EpisodicEvent.event_type == "document_ingested",
                        )
                    )
                    is not None
                )

            assert parser.calls == 1
            assert provider.calls >= 1
    finally:
        await _clean_owner(database, owner_id)
        await parser.aclose()
        await embeddings.aclose()
        await inference.aclose()
        await database.close()


@pytest.mark.asyncio
@pytest.mark.parametrize("tampered_field", ["content", "received_at"])
async def test_staged_integrity_tampering_terminally_fails_and_removes_stage(
    tampered_field: str,
) -> None:
    owner_id = f"document-{tampered_field}-tamper-{uuid4().hex}"
    token = uuid4().hex
    settings = _settings(owner_id, token)
    database = Database(settings)
    inference = DeterministicInference()
    provider = ProbedEmbeddingProvider()
    embeddings = _embedding_service(provider)
    parser = ProbedDocumentParser()
    application = create_app(
        settings=settings,
        database=database,
        inference=inference,
        embeddings=embeddings,
    )
    headers = {"Authorization": f"Bearer {token}"}
    content = b"The approved bytes must remain immutable."

    try:
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=application),
            base_url="http://testserver",
        ) as client:
            created = await _upload_txt(client, headers=headers, content=content)
            task_id = UUID(created["id"])
            async with database.session_factory() as session:
                task = await session.get(TaskQueue, task_id)
                assert task is not None
                staging_id = UUID(task.payload["staging_id"])

            approved = await client.post(
                f"/v1/tasks/{task_id}/approve",
                headers=headers,
                json={"action_digest": created["action_digest"]},
            )
            assert approved.status_code == 200

            async with database.session_factory() as session, session.begin():
                replacement: object
                if tampered_field == "content":
                    replacement = bytes([content[0] ^ 1]) + content[1:]
                else:
                    staged_received_at = await session.scalar(
                        select(DocumentStaging.received_at).where(DocumentStaging.id == staging_id)
                    )
                    assert staged_received_at is not None
                    replacement = staged_received_at + timedelta(seconds=1)
                await session.execute(
                    update(DocumentStaging)
                    .where(DocumentStaging.id == staging_id)
                    .values({tampered_field: replacement})
                )

            worker = Worker(
                settings=settings,
                database=database,
                inference=inference,
                embeddings=embeddings,
                document_parser=cast(DocumentParser, parser),
            )
            assert await worker.run_once() is True

            async with database.session_factory() as session:
                task = await session.get(TaskQueue, task_id)
                assert task is not None
                assert task.status == "failed"
                assert task.error_text == "approved document metadata does not match staged content"
                assert await session.get(DocumentStaging, staging_id) is None
                event_types = list(
                    (
                        await session.scalars(
                            select(EpisodicEvent.event_type).where(
                                EpisodicEvent.owner_id == owner_id,
                                EpisodicEvent.trace_id == task.trace_id,
                            )
                        )
                    ).all()
                )
                assert "document_ingest_failed" in event_types
                assert "task_failed" in event_types

            assert parser.calls == 0
            assert provider.calls == 0
    finally:
        await _clean_owner(database, owner_id)
        await embeddings.aclose()
        await inference.aclose()
        await database.close()


@pytest.mark.asyncio
async def test_retryable_parser_timeout_requeues_then_succeeds_with_staging_intact() -> None:
    owner_id = f"document-parser-retry-{uuid4().hex}"
    token = uuid4().hex
    settings = _settings(owner_id, token)
    database = Database(settings)
    inference = DeterministicInference()
    provider = ProbedEmbeddingProvider()
    embeddings = _embedding_service(provider)
    parser = FlakyDocumentParser()
    application = create_app(
        settings=settings,
        database=database,
        inference=inference,
        embeddings=embeddings,
    )
    headers = {"Authorization": f"Bearer {token}"}

    try:
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=application),
            base_url="http://testserver",
        ) as client:
            created = await _upload_txt(
                client,
                headers=headers,
                content=b"Retry this document after parser infrastructure recovers.",
            )
            task_id = UUID(created["id"])
            async with database.session_factory() as session:
                task = await session.get(TaskQueue, task_id)
                assert task is not None
                staging_id = UUID(task.payload["staging_id"])

            approved = await client.post(
                f"/v1/tasks/{task_id}/approve",
                headers=headers,
                json={"action_digest": created["action_digest"]},
            )
            assert approved.status_code == 200

            worker = Worker(
                settings=settings,
                database=database,
                inference=inference,
                embeddings=embeddings,
                document_parser=cast(DocumentParser, parser),
            )
            assert await worker.run_once() is True

            async with database.session_factory() as session:
                task = await session.get(TaskQueue, task_id)
                assert task is not None
                assert task.status == "queued"
                assert task.attempt_count == 1
                assert task.error_text == "parser_timeout"
                assert await session.get(DocumentStaging, staging_id) is not None
                event_types = list(
                    (
                        await session.scalars(
                            select(EpisodicEvent.event_type).where(
                                EpisodicEvent.owner_id == owner_id,
                                EpisodicEvent.trace_id == task.trace_id,
                            )
                        )
                    ).all()
                )
                assert "task_requeued" in event_types
                assert "document_ingest_failed" not in event_types

            # Bypass the deterministic retry backoff without sleeping; the second
            # claim still exercises approval reuse and a fresh execution token.
            async with database.session_factory() as session, session.begin():
                await session.execute(
                    update(TaskQueue)
                    .where(TaskQueue.id == task_id)
                    .values(run_after=datetime.now(UTC) - timedelta(seconds=1))
                )

            assert await worker.run_once() is True

            async with database.session_factory() as session:
                task = await session.get(TaskQueue, task_id)
                assert task is not None
                assert task.status == "done"
                assert task.attempt_count == 2
                assert task.result is not None and task.result["created"] is True
                assert await session.get(DocumentStaging, staging_id) is None
                event_types = list(
                    (
                        await session.scalars(
                            select(EpisodicEvent.event_type).where(
                                EpisodicEvent.owner_id == owner_id,
                                EpisodicEvent.trace_id == task.trace_id,
                            )
                        )
                    ).all()
                )
                assert event_types.count("task_requeued") == 1
                assert event_types.count("document_ingested") == 1
                assert event_types.count("task_completed") == 1

            assert parser.calls == 2
            assert provider.calls == 1
    finally:
        await _clean_owner(database, owner_id)
        await parser.aclose()
        await embeddings.aclose()
        await inference.aclose()
        await database.close()


@pytest.mark.asyncio
async def test_slow_parser_receives_lease_heartbeats_until_completion() -> None:
    owner_id = f"document-heartbeat-{uuid4().hex}"
    token = uuid4().hex
    settings = _settings(owner_id, token)
    database = Database(settings)
    inference = DeterministicInference()
    provider = ProbedEmbeddingProvider()
    embeddings = _embedding_service(provider)
    parser = GatedDocumentParser()
    application = create_app(
        settings=settings,
        database=database,
        inference=inference,
        embeddings=embeddings,
    )
    headers = {"Authorization": f"Bearer {token}"}
    worker_task: asyncio.Task[bool] | None = None

    try:
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=application),
            base_url="http://testserver",
        ) as client:
            created = await _upload_txt(
                client,
                headers=headers,
                content=b"Hold parsing long enough to require a lease heartbeat.",
            )
            task_id = UUID(created["id"])
            approved = await client.post(
                f"/v1/tasks/{task_id}/approve",
                headers=headers,
                json={"action_digest": created["action_digest"]},
            )
            assert approved.status_code == 200

            worker = Worker(
                settings=settings,
                database=database,
                inference=inference,
                embeddings=embeddings,
                document_parser=cast(DocumentParser, parser),
            )
            worker_task = asyncio.create_task(worker.run_once())
            await asyncio.wait_for(parser.started.wait(), timeout=2)

            async with database.session_factory() as session:
                task = await session.get(TaskQueue, task_id)
                assert task is not None
                assert task.status == "running"
                initial_expiry = task.lease_expires_at
                assert initial_expiry is not None

            renewed_expiry: datetime | None = None
            for _attempt in range(24):
                await asyncio.sleep(0.25)
                async with database.session_factory() as session:
                    current_expiry = await session.scalar(
                        select(TaskQueue.lease_expires_at).where(TaskQueue.id == task_id)
                    )
                if current_expiry is not None and current_expiry > initial_expiry:
                    renewed_expiry = current_expiry
                    break
            assert renewed_expiry is not None

            parser.release.set()
            assert await asyncio.wait_for(worker_task, timeout=5) is True

            async with database.session_factory() as session:
                task = await session.get(TaskQueue, task_id)
                assert task is not None
                assert task.status == "done"
                assert task.attempt_count == 1
                assert task.lease_expires_at is None
                assert task.result is not None and task.result["created"] is True

            assert parser.calls == 1
            assert provider.calls == 1
    finally:
        parser.release.set()
        if worker_task is not None and not worker_task.done():
            worker_task.cancel()
            with suppress(asyncio.CancelledError):
                await worker_task
        await _clean_owner(database, owner_id)
        await parser.aclose()
        await embeddings.aclose()
        await inference.aclose()
        await database.close()


@pytest.mark.asyncio
async def test_document_cancellation_deletes_staging_object() -> None:
    owner_id = f"document-cancel-{uuid4().hex}"
    token = uuid4().hex
    settings = _settings(owner_id, token)
    database = Database(settings)
    inference = DeterministicInference()
    embeddings = _embedding_service(ProbedEmbeddingProvider())
    application = create_app(
        settings=settings,
        database=database,
        inference=inference,
        embeddings=embeddings,
    )
    headers = {"Authorization": f"Bearer {token}"}

    try:
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=application),
            base_url="http://testserver",
        ) as client:
            created = await _upload_txt(
                client,
                headers=headers,
                content=b"Cancel this staged document.",
            )
            task_id = UUID(created["id"])
            async with database.session_factory() as session:
                task = await session.get(TaskQueue, task_id)
                assert task is not None
                staging_id = UUID(task.payload["staging_id"])

            cancelled = await client.post(f"/v1/tasks/{task_id}/cancel", headers=headers)
            assert cancelled.status_code == 204

            async with database.session_factory() as session:
                task = await session.get(TaskQueue, task_id)
                assert task is not None and task.status == "cancelled"
                assert await session.get(DocumentStaging, staging_id) is None
                assert (
                    await session.scalar(
                        select(EpisodicEvent.id).where(
                            EpisodicEvent.owner_id == owner_id,
                            EpisodicEvent.trace_id == task.trace_id,
                            EpisodicEvent.event_type == "task_cancelled",
                        )
                    )
                    is not None
                )
    finally:
        await _clean_owner(database, owner_id)
        await embeddings.aclose()
        await inference.aclose()
        await database.close()


@pytest.mark.asyncio
async def test_expired_unapproved_document_is_cancelled_and_cleaned_by_worker() -> None:
    owner_id = f"document-expiry-{uuid4().hex}"
    token = uuid4().hex
    settings = _settings(owner_id, token)
    database = Database(settings)
    inference = DeterministicInference()
    provider = ProbedEmbeddingProvider()
    embeddings = _embedding_service(provider)
    parser = ProbedDocumentParser()
    application = create_app(
        settings=settings,
        database=database,
        inference=inference,
        embeddings=embeddings,
    )
    headers = {"Authorization": f"Bearer {token}"}

    try:
        async with httpx.AsyncClient(
            transport=httpx.ASGITransport(app=application),
            base_url="http://testserver",
        ) as client:
            created = await _upload_txt(
                client,
                headers=headers,
                content=b"This approval will expire before parsing.",
            )
            task_id = UUID(created["id"])
            async with database.session_factory() as session:
                task = await session.get(TaskQueue, task_id)
                assert task is not None
                staging_id = UUID(task.payload["staging_id"])

            expired_at = datetime.now(UTC) - timedelta(seconds=1)
            async with database.session_factory() as session, session.begin():
                await session.execute(
                    update(DocumentStaging)
                    .where(DocumentStaging.id == staging_id)
                    .values(expires_at=expired_at)
                )
                await session.execute(
                    update(TaskQueue)
                    .where(TaskQueue.id == task_id)
                    .values(approval_expires_at=expired_at)
                )

            worker = Worker(
                settings=settings,
                database=database,
                inference=inference,
                embeddings=embeddings,
                document_parser=cast(DocumentParser, parser),
            )
            assert await worker.run_once() is False

            async with database.session_factory() as session:
                task = await session.get(TaskQueue, task_id)
                assert task is not None
                assert task.status == "cancelled"
                assert task.error_text == "document upload approval expired"
                assert await session.get(DocumentStaging, staging_id) is None
                failed_event = await session.scalar(
                    select(EpisodicEvent).where(
                        EpisodicEvent.owner_id == owner_id,
                        EpisodicEvent.trace_id == task.trace_id,
                        EpisodicEvent.event_type == "document_ingest_failed",
                    )
                )
                assert failed_event is not None
                assert failed_event.payload["error_code"] == "approval_expired"

            assert parser.calls == 0
            assert provider.calls == 0
    finally:
        await _clean_owner(database, owner_id)
        await embeddings.aclose()
        await inference.aclose()
        await database.close()


@pytest.mark.asyncio
async def test_database_rejects_cross_owner_staging_task_binding() -> None:
    owner_id = f"document-owner-a-{uuid4().hex}"
    foreign_owner = f"document-owner-b-{uuid4().hex}"
    settings = _settings(owner_id, uuid4().hex)
    database = Database(settings)
    task_id: UUID

    try:
        async with database.session_factory() as session, session.begin():
            task = await TaskRepository(session).create(
                owner_id=owner_id,
                kind="document.ingest",
                risk_level="local_mutation",
                status="waiting_approval",
                trace_id=f"trc_{uuid4().hex}",
                payload={
                    "staging_id": str(uuid4()),
                    "original_filename": "notes.txt",
                    "source_sha256": hashlib.sha256(b"owner scoped").hexdigest(),
                    "detected_mime_type": "text/plain",
                    "byte_size": len(b"owner scoped"),
                    "source_timestamp": "2026-07-22T12:30:00Z",
                    "received_at": "2026-07-22T12:31:00Z",
                    "source_time_basis": "user_supplied",
                    "title": None,
                    "sensitivity": "private",
                    "retention_class": "keep",
                },
                action_digest="a" * 64,
                approval_expires_at=datetime.now(UTC) + timedelta(minutes=5),
            )
            task_id = task.id

        with pytest.raises(IntegrityError):
            async with database.session_factory() as session, session.begin():
                session.add(
                    DocumentStaging(
                        id=uuid4(),
                        owner_id=foreign_owner,
                        task_id=task_id,
                        original_filename="notes.txt",
                        detected_mime_type="text/plain",
                        source_sha256=hashlib.sha256(b"owner scoped").digest(),
                        byte_size=len(b"owner scoped"),
                        content=b"owner scoped",
                        source_timestamp=datetime(2026, 7, 22, 12, 30, tzinfo=UTC),
                        received_at=datetime(2026, 7, 22, 12, 31, tzinfo=UTC),
                        expires_at=datetime.now(UTC) + timedelta(minutes=5),
                    )
                )
                await session.flush()
    finally:
        await _clean_owner(database, owner_id)
        await _clean_owner(database, foreign_owner)
        await database.close()
