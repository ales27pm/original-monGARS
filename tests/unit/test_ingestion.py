from __future__ import annotations

import stat
from dataclasses import replace
from datetime import UTC, datetime
from io import BytesIO
from uuid import UUID
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import pytest
from docx import Document
from pypdf import PdfWriter

from mongars.ingestion import (
    ContentTypeMismatchError,
    DocumentIngestionService,
    DocumentLimits,
    DocumentLocator,
    DocumentStructureLimitError,
    DocumentTooLargeError,
    EncryptedDocumentError,
    ExtractedSegment,
    ExtractedTextTooLargeError,
    IngestionContext,
    InvalidFilenameError,
    IsolatedDocumentParser,
    MalformedDocumentError,
    NoUsableTextError,
    ParserProcessLimits,
    ParserTimeoutError,
    UnsafeDocumentError,
    UnsupportedDocumentTypeError,
    UploadEnvelope,
    chunk_segments,
)
from mongars.ingestion.extractors.text import TextExtractor
from mongars.ingestion.registry import ExtractorRegistry

_TASK_ID = UUID("018f0a1e-8d46-7a3e-83c6-75ba32f83601")


def _context() -> IngestionContext:
    return IngestionContext(
        owner_id="owner-a",
        ingestion_task_id=_TASK_ID,
        sensitivity="restricted",
        retention_class="ttl_30d",
        source_timestamp=datetime(2026, 7, 22, 12, 30, tzinfo=UTC),
        received_at=datetime(2026, 7, 22, 12, 31, tzinfo=UTC),
    )


def _extract(
    filename: str,
    mime_type: str,
    content: bytes,
    *,
    limits: DocumentLimits | None = None,
):
    service = DocumentIngestionService(limits=limits)
    validated = service.validate_envelope(
        UploadEnvelope(
            original_filename=filename,
            declared_mime_type=mime_type,
            declared_size=len(content),
            content=content,
        )
    )
    return service.extract_validated(validated, context=_context())


def _make_pdf(text: str = "Hello from PDF") -> bytes:
    objects: list[bytes] = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>"
        ),
    ]
    stream = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode()
    objects.extend(
        (
            b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
            b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        )
    )
    output = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for index, value in enumerate(objects, 1):
        offsets.append(len(output))
        output.extend(f"{index} 0 obj\n".encode())
        output.extend(value)
        output.extend(b"\nendobj\n")
    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets:
        output.extend(f"{offset:010d} 00000 n \n".encode())
    output.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode()
    )
    return bytes(output)


def _make_docx() -> bytes:
    document = Document()
    document.add_heading("Safety report", level=1)
    document.add_paragraph("A bounded DOCX body.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "monGARS"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _rewrite_docx(
    content: bytes,
    *,
    replacements: dict[str, bytes] | None = None,
    additions: tuple[tuple[ZipInfo | str, bytes], ...] = (),
) -> bytes:
    replacements = replacements or {}
    output = BytesIO()
    with ZipFile(BytesIO(content)) as source, ZipFile(output, "w", ZIP_DEFLATED) as target:
        for member in source.infolist():
            value = replacements.get(member.filename, source.read(member.filename))
            target.writestr(member, value)
        for name, value in additions:
            target.writestr(name, value)
    return output.getvalue()


def test_txt_extraction_returns_bounded_provenance() -> None:
    result = _extract("notes.txt", "text/plain; charset=utf-8", b"First\r\n\r\nSecond\n")

    assert result.text == "First\n\nSecond"
    assert result.provenance.owner_id == "owner-a"
    assert result.provenance.ingestion_task_id == _TASK_ID
    assert result.provenance.validated_mime_type == "text/plain"
    assert result.provenance.byte_size == 16
    assert result.provenance.extracted_character_count == len(result.text)
    assert result.provenance.section_count == 2
    assert result.provenance.parser_name == "utf8-text"
    assert result.provenance.as_metadata()["source_timestamp"] == "2026-07-22T12:30:00+00:00"
    assert result.provenance.as_metadata()["received_at"] == "2026-07-22T12:31:00+00:00"
    assert result.provenance.as_metadata()["source_time_basis"] == "user_supplied"
    assert [segment.text for segment in result.segments] == ["First", "Second"]
    assert [
        (segment.locator.line_start, segment.locator.line_end) for segment in result.segments
    ] == [(1, 1), (3, 3)]


def test_markdown_extraction_preserves_source_semantics() -> None:
    result = _extract(
        "design.MD",
        "text/x-markdown",
        b"# Cortex\n\n- bounded context\n- explicit policy\n",
    )

    assert result.text.startswith("# Cortex")
    assert "- explicit policy" in result.text
    assert result.provenance.validated_mime_type == "text/markdown"
    assert result.segments[0].locator.heading_path == ("Cortex",)
    assert result.segments[1].locator.heading_path == ("Cortex",)


def test_markdown_tracks_nested_headings_without_trusting_fenced_code() -> None:
    result = _extract(
        "guide.md",
        "text/markdown",
        b"""# Root
Intro
## Child
Details
```markdown
# Not a heading
```
After code
# Next
Done
""",
    )

    assert [segment.locator.heading_path for segment in result.segments] == [
        ("Root",),
        ("Root",),
        ("Root", "Child"),
        ("Root", "Child"),
        ("Root", "Child"),
        ("Root", "Child"),
        ("Next",),
        ("Next",),
    ]
    assert result.segments[4].text.startswith("```markdown")
    assert result.segments[4].locator.line_start == 5
    assert result.segments[4].locator.line_end == 7


def test_markdown_tracks_setext_headings_as_major_boundaries() -> None:
    result = _extract(
        "setext.md",
        "text/markdown",
        b"Root title\n==========\nIntro\n\nChild title\n-----------\nDetails\n",
    )

    assert [segment.locator.heading_path for segment in result.segments] == [
        ("Root title",),
        ("Root title",),
        ("Root title", "Child title"),
        ("Root title", "Child title"),
    ]
    assert result.segments[0].locator.line_end == 2
    assert result.segments[2].locator.line_start == 5


def test_markdown_does_not_strip_hash_from_heading_text() -> None:
    result = _extract(
        "language.md",
        "text/markdown",
        b"# C#\nCode notes\n",
    )

    assert result.segments[0].locator.heading_path == ("C#",)
    assert result.segments[1].locator.heading_path == ("C#",)


@pytest.mark.parametrize("marker", ["```oops", "~~~oops"])
def test_markdown_rejects_false_fence_closer_as_heading_boundary(marker: str) -> None:
    fence = marker[0] * 3
    result = _extract(
        "fence.md",
        "text/markdown",
        f"{fence}\n{marker}\n# Still code\n{fence}\nAfter\n".encode(),
    )

    assert len(result.segments) == 2
    assert result.segments[0].locator.heading_path == ()
    assert "# Still code" in result.segments[0].text
    assert result.segments[1].locator.heading_path == ()


def test_heading_metadata_limits_fail_with_typed_structure_error() -> None:
    heading = "x" * 501
    with pytest.raises(DocumentStructureLimitError, match="heading metadata"):
        _extract("long.md", "text/markdown", f"# {heading}\nbody".encode())
    with pytest.raises(DocumentStructureLimitError, match="heading metadata"):
        _extract(
            "long.html",
            "text/html",
            f"<!doctype html><html><body><h1>{heading}</h1></body></html>".encode(),
        )

    document = Document()
    document.add_heading(heading, level=1)
    output = BytesIO()
    document.save(output)
    with pytest.raises(DocumentStructureLimitError, match="heading metadata"):
        _extract(
            "long.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            output.getvalue(),
        )


def test_html_sanitization_preserves_visible_text_and_removes_known_hidden_text() -> None:
    result = _extract(
        "page.html",
        "text/html",
        b"""<!doctype html><html><head><title>Not body content</title></head><body>
        <!-- secret --><main><h1>Visible</h1><p hidden>Hidden one</p>
        <p style="display: none">Hidden two</p><p>Useful body</p></main></body></html>""",
    )

    assert "Visible" in result.text
    assert "Useful body" in result.text
    assert "secret" not in result.text
    assert "Hidden" not in result.text
    assert "Not body content" not in result.text
    assert result.segments[0].locator.heading_path == ("Visible",)
    assert result.segments[1].locator.heading_path == ("Visible",)


def test_html_preserves_direct_text_heading_hierarchy_and_table_cells() -> None:
    result = _extract(
        "structured.html",
        "text/html",
        b"""<!doctype html><html><body>Intro <strong>bold</strong>
        <h1>Install</h1>Before table
        <table><tr><th>Name</th><th>Value</th></tr>
        <tr><td>GPU</td><td>RTX 2070</td></tr></table></body></html>""",
    )

    assert result.segments[0].text == "Intro bold"
    assert result.segments[1].locator.heading_path == ("Install",)
    assert result.segments[2].text == "Before table"
    table_segments = result.segments[3:]
    assert [segment.locator.cell_reference for segment in table_segments] == [
        "A1",
        "B1",
        "A2",
        "B2",
    ]
    assert all(segment.locator.table_index == 0 for segment in table_segments)
    assert all(segment.locator.heading_path == ("Install",) for segment in table_segments)


def test_html_table_coordinates_respect_rowspan_and_colspan() -> None:
    result = _extract(
        "spans.html",
        "text/html",
        b"""<!doctype html><html><body><table>
        <tr><th rowspan="2">Name</th><th colspan="2">Values</th></tr>
        <tr><td>Primary</td><td>Secondary</td></tr>
        </table></body></html>""",
    )

    assert [segment.locator.cell_reference for segment in result.segments] == [
        "A1",
        "B1",
        "B2",
        "C2",
    ]


def test_html_rejects_stylesheet_hidden_content_ambiguity() -> None:
    content = (
        b"<!doctype html><html><head><style>.hidden{display:none}</style></head>"
        b"<body><p class=hidden>IGNORE</p></body></html>"
    )

    with pytest.raises(UnsafeDocumentError, match="stylesheet"):
        _extract("page.html", "text/html", content)


@pytest.mark.parametrize(
    "fragment",
    [
        '<link rel="stylesheet" href="theme.css"><p>text</p>',
        '<p class="possibly-hidden">text</p>',
        '<p id="possibly-hidden">text</p>',
    ],
)
def test_html_rejects_css_selector_ambiguity(fragment: str) -> None:
    content = f"<!doctype html><html><body>{fragment}</body></html>".encode()
    with pytest.raises(UnsafeDocumentError, match=r"stylesheet|selectors"):
        _extract("page.html", "text/html", content)


@pytest.mark.parametrize(
    "fragment",
    [
        '<script src="https://evil.invalid/x.js"></script>',
        '<form action="https://evil.invalid"><input></form>',
        '<iframe src="https://evil.invalid"></iframe>',
        '<svg><a href="https://evil.invalid">x</a></svg>',
        "<math><mi>x</mi></math>",
        '<p onclick="steal()">click</p>',
        '<a href="javascript:steal()">click</a>',
        '<a href="data:text/html,attack">click</a>',
        '<meta http-equiv="refresh" content="0;https://evil.invalid">',
        '<p style="color: red">styled</p>',
    ],
)
def test_html_rejects_active_content(fragment: str) -> None:
    content = f"<!doctype html><html><body><p>safe</p>{fragment}</body></html>".encode()

    with pytest.raises(UnsafeDocumentError):
        _extract("page.html", "text/html", content)


def test_html_rejects_external_doctype() -> None:
    content = b'<!DOCTYPE html SYSTEM "file:///etc/passwd"><html><body>safe</body></html>'

    with pytest.raises(UnsafeDocumentError, match="document type"):
        _extract("page.html", "text/html", content)


def test_pdf_extraction_and_page_provenance() -> None:
    result = _extract("report.pdf", "application/pdf", _make_pdf())

    assert result.text == "Hello from PDF"
    assert result.provenance.page_count == 1
    assert result.provenance.section_count == 1
    assert result.provenance.parser_name == "pypdf"
    assert result.segments[0].locator.page_number == 1
    assert result.segments[0].locator.block_index == 0


def test_pdf_segments_never_cross_page_boundaries() -> None:
    writer = PdfWriter()
    writer.append(BytesIO(_make_pdf("First page")))
    writer.append(BytesIO(_make_pdf("Second page")))
    output = BytesIO()
    writer.write(output)

    result = _extract("pages.pdf", "application/pdf", output.getvalue())

    assert [segment.text for segment in result.segments] == ["First page", "Second page"]
    assert [segment.locator.page_number for segment in result.segments] == [1, 2]
    assert [segment.locator.block_index for segment in result.segments] == [0, 0]


def test_pdf_rejects_encryption_even_when_password_is_empty() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.encrypt("")
    output = BytesIO()
    writer.write(output)

    with pytest.raises(EncryptedDocumentError):
        _extract("secret.pdf", "application/pdf", output.getvalue())


def test_pdf_rejects_excessive_pages_before_extraction() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.add_blank_page(width=100, height=100)
    output = BytesIO()
    writer.write(output)

    with pytest.raises(DocumentStructureLimitError, match="page limit"):
        _extract(
            "pages.pdf",
            "application/pdf",
            output.getvalue(),
            limits=DocumentLimits(max_pages=1),
        )


def test_pdf_without_text_is_rejected() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    output = BytesIO()
    writer.write(output)

    with pytest.raises(NoUsableTextError):
        _extract("blank.pdf", "application/pdf", output.getvalue())


def test_docx_extracts_body_and_table() -> None:
    result = _extract(
        "report.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        _make_docx(),
    )

    assert "Safety report" in result.text
    assert "A bounded DOCX body." in result.text
    assert "Name" in result.text
    assert "monGARS" in result.text
    assert result.provenance.parser_name == "python-docx"
    assert result.provenance.section_count == 3
    assert result.segments[0].locator.heading_path == ("Safety report",)
    table_segments = [
        segment for segment in result.segments if segment.locator.table_index is not None
    ]
    assert [segment.locator.cell_reference for segment in table_segments] == ["A1", "B1"]
    assert all(segment.locator.heading_path == ("Safety report",) for segment in table_segments)


def test_docx_merged_cell_is_emitted_once_at_its_top_left_coordinate() -> None:
    document = Document()
    table = document.add_table(rows=1, cols=2)
    merged = table.cell(0, 0).merge(table.cell(0, 1))
    merged.text = "Merged value"
    output = BytesIO()
    document.save(output)

    result = _extract(
        "merged.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        output.getvalue(),
    )

    assert [segment.text for segment in result.segments] == ["Merged value"]
    assert result.segments[0].locator.cell_reference == "A1"


def test_docx_rejects_external_relationship() -> None:
    original = _make_docx()
    with ZipFile(BytesIO(original)) as archive:
        relationships = archive.read("word/_rels/document.xml.rels")
    injected = relationships.replace(
        b"</Relationships>",
        (
            b'<Relationship Id="evil" Type="urn:evil" '
            b'Target="file:///etc/passwd" TargetMode="External"/></Relationships>'
        ),
    )
    malicious = _rewrite_docx(
        original,
        replacements={"word/_rels/document.xml.rels": injected},
    )

    with pytest.raises(UnsafeDocumentError, match="external relationship"):
        _extract(
            "external.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            malicious,
        )


@pytest.mark.parametrize("member_name", ["../escape.txt", "/absolute.txt", "C:/host.txt"])
def test_docx_rejects_unsafe_archive_paths(member_name: str) -> None:
    malicious = _rewrite_docx(_make_docx(), additions=((member_name, b"attack"),))

    with pytest.raises(UnsafeDocumentError):
        _extract(
            "unsafe.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            malicious,
        )


def test_docx_rejects_symbolic_link_member() -> None:
    symlink = ZipInfo("word/link")
    symlink.create_system = 3
    symlink.external_attr = (stat.S_IFLNK | 0o777) << 16
    malicious = _rewrite_docx(_make_docx(), additions=((symlink, b"/etc/passwd"),))

    with pytest.raises(UnsafeDocumentError, match="symbolic link"):
        _extract(
            "unsafe.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            malicious,
        )


def test_docx_rejects_decompression_bomb_metadata() -> None:
    malicious = _rewrite_docx(
        _make_docx(),
        additions=(("word/bomb.txt", b"A" * 1_000_000),),
    )

    with pytest.raises(DocumentStructureLimitError, match="compression ratio"):
        _extract(
            "bomb.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            malicious,
            limits=DocumentLimits(max_compression_ratio=100),
        )


def test_docx_rejects_truncated_archive() -> None:
    truncated = _make_docx()[:-32]

    with pytest.raises(MalformedDocumentError, match="malformed or truncated"):
        _extract(
            "truncated.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            truncated,
        )


@pytest.mark.parametrize(
    "filename",
    [
        "../note.txt",
        "/note.txt",
        "folder\\note.txt",
        "bad\x00.txt",
        "invoice\u202efdp.txt",
        "invisible\u200b.txt",
    ],
)
def test_filename_must_be_a_safe_basename(filename: str) -> None:
    service = DocumentIngestionService()

    with pytest.raises(InvalidFilenameError):
        service.validate_envelope(UploadEnvelope(filename, "text/plain", b"safe"))


def test_filename_is_nfc_normalized_and_confusable_separators_are_rejected() -> None:
    service = DocumentIngestionService()
    validated = service.validate_envelope(
        UploadEnvelope("re\u0301sume\u0301.txt", "text/plain", b"safe")
    )

    assert validated.original_filename == "résumé.txt"
    with pytest.raises(InvalidFilenameError):
        service.validate_envelope(UploadEnvelope("folder\u2215note.txt", "text/plain", b"safe"))


def test_unsupported_extension_and_mime_are_rejected() -> None:
    service = DocumentIngestionService()

    with pytest.raises(UnsupportedDocumentTypeError):
        service.validate_envelope(UploadEnvelope("archive.zip", "application/zip", b"PK"))
    with pytest.raises(UnsupportedDocumentTypeError):
        service.validate_envelope(UploadEnvelope("note.txt", "application/octet-stream", b"safe"))


@pytest.mark.parametrize(
    ("filename", "mime_type", "content"),
    [
        ("report.pdf", "text/plain", _make_pdf()),
        ("report.txt", "text/plain", _make_pdf()),
        ("report.pdf", "application/pdf", b"plain UTF-8 text"),
        ("page.html", "text/html", b"plain UTF-8 text"),
        ("page.txt", "text/plain", b"<html><body>HTML</body></html>"),
    ],
)
def test_mime_extension_and_bytes_must_agree(
    filename: str,
    mime_type: str,
    content: bytes,
) -> None:
    with pytest.raises(ContentTypeMismatchError):
        _extract(filename, mime_type, content)


def test_declared_and_streamed_size_must_agree() -> None:
    service = DocumentIngestionService()

    with pytest.raises(MalformedDocumentError, match="size"):
        service.validate_envelope(
            UploadEnvelope("note.txt", "text/plain", b"hello", declared_size=500)
        )


def test_input_and_extracted_character_limits_are_independent() -> None:
    with pytest.raises(DocumentTooLargeError):
        _extract(
            "note.txt",
            "text/plain",
            b"0123456789",
            limits=DocumentLimits(max_input_bytes=5),
        )
    with pytest.raises(ExtractedTextTooLargeError):
        _extract(
            "note.txt",
            "text/plain",
            b"0123456789",
            limits=DocumentLimits(max_extracted_chars=5),
        )


def test_empty_binary_and_excessive_sections_are_rejected() -> None:
    with pytest.raises(NoUsableTextError):
        _extract("empty.txt", "text/plain", b"  \n\n ")
    with pytest.raises(MalformedDocumentError, match="control"):
        _extract("binary.txt", "text/plain", b"safe\x01unsafe")
    with pytest.raises(DocumentStructureLimitError, match="section"):
        _extract(
            "sections.txt",
            "text/plain",
            b"one\n\ntwo",
            limits=DocumentLimits(max_sections=1),
        )


def test_staged_bytes_are_revalidated_before_parser_dispatch() -> None:
    service = DocumentIngestionService()
    validated = service.validate_envelope(UploadEnvelope("note.txt", "text/plain", b"original"))
    tampered = replace(validated, content=b"attacker")

    with pytest.raises(UnsafeDocumentError, match="integrity"):
        service.extract_validated(tampered, context=_context())


def test_context_requires_governance_and_timezone() -> None:
    service = DocumentIngestionService()
    validated = service.validate_envelope(UploadEnvelope("note.txt", "text/plain", b"safe"))

    with pytest.raises(ValueError, match="sensitivity"):
        service.extract_validated(validated, context=replace(_context(), sensitivity="public"))
    with pytest.raises(ValueError, match="timezone-aware"):
        service.extract_validated(
            validated,
            context=replace(_context(), source_timestamp=datetime(2026, 7, 22)),
        )


def test_registry_rejects_duplicate_media_type() -> None:
    with pytest.raises(ValueError, match="already registered"):
        ExtractorRegistry((TextExtractor(), TextExtractor()))


def test_locator_preserving_chunker_never_crosses_structural_boundaries() -> None:
    first_locator = DocumentLocator(
        media_type="application/pdf",
        page_number=1,
        block_index=0,
    )
    second_locator = DocumentLocator(
        media_type="application/pdf",
        page_number=2,
        block_index=0,
    )
    segments = (
        ExtractedSegment(
            text=" ".join(f"first-{index}" for index in range(40)),
            locator=first_locator,
        ),
        ExtractedSegment(
            text=" ".join(f"second-{index}" for index in range(40)),
            locator=second_locator,
        ),
    )

    chunks = chunk_segments(
        segments,
        max_tokens=32,
        overlap_tokens=4,
        max_characters=1_000,
    )

    assert len(chunks) == 4
    assert all(chunk.locator == first_locator for chunk in chunks[:2])
    assert all(chunk.locator == second_locator for chunk in chunks[2:])
    assert all("second-" not in chunk.text for chunk in chunks[:2])
    assert all("first-" not in chunk.text for chunk in chunks[2:])


def test_locator_preserving_chunker_narrows_multiline_source_ranges() -> None:
    segment = ExtractedSegment(
        text=(
            " ".join(f"line-one-{index}" for index in range(20))
            + "\n"
            + " ".join(f"line-two-{index}" for index in range(20))
        ),
        locator=DocumentLocator(
            media_type="text/plain",
            block_index=0,
            line_start=10,
            line_end=11,
        ),
    )

    chunks = chunk_segments(
        (segment,),
        max_tokens=32,
        overlap_tokens=4,
        max_characters=1_000,
    )

    assert [chunk.locator.line_start for chunk in chunks] == [10, 11]
    assert [chunk.locator.line_end for chunk in chunks] == [10, 11]


def test_document_locator_rejects_ambiguous_or_invalid_coordinates() -> None:
    with pytest.raises(ValueError, match="both endpoints"):
        DocumentLocator(
            media_type="text/plain",
            block_index=0,
            line_start=1,
        )
    with pytest.raises(ValueError, match="requires a table"):
        DocumentLocator(
            media_type="text/plain",
            block_index=0,
            cell_reference="A1",
        )
    with pytest.raises(ValueError, match="page number"):
        DocumentLocator(
            media_type="application/pdf",
            block_index=0,
            page_number=0,
        )


@pytest.mark.asyncio
async def test_isolated_parser_returns_result_from_disposable_subprocess() -> None:
    service = DocumentIngestionService()
    validated = service.validate_envelope(UploadEnvelope("note.txt", "text/plain", b"isolated"))

    result = await IsolatedDocumentParser().extract(validated)

    assert result.text == "isolated"
    assert result.parser_name == "utf8-text"


@pytest.mark.asyncio
async def test_isolated_parser_enforces_wall_clock_timeout() -> None:
    service = DocumentIngestionService()
    validated = service.validate_envelope(UploadEnvelope("note.txt", "text/plain", b"isolated"))
    parser = IsolatedDocumentParser(process_limits=ParserProcessLimits(timeout_seconds=0.000_001))

    with pytest.raises(ParserTimeoutError):
        await parser.extract(validated)
